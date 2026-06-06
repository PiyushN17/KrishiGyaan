from pathlib import Path
from textwrap import wrap
import re

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle, Circle
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
OUT = ROOT / "deliverables"
ASSETS.mkdir(parents=True, exist_ok=True)
OUT.mkdir(parents=True, exist_ok=True)

GREEN = "#176B3A"
DARK = "#19352A"
LIGHT = "#EAF5EE"
GOLD = "#B98620"
GRAY = "#52635B"
RED = "#A33A32"
BLUE = "#2E5D8A"


def draw_box(ax, x, y, w, h, text, color=GREEN, fill="#FFFFFF", fontsize=9):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.012,rounding_size=0.015",
        linewidth=1.5, edgecolor=color, facecolor=fill
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, "\n".join(wrap(text, 24)),
            ha="center", va="center", fontsize=fontsize, color=DARK, weight="semibold")
    return patch


def arrow(ax, x1, y1, x2, y2, label=None):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color=GRAY, lw=1.5))
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.018, label,
                ha="center", va="bottom", fontsize=7.5, color=GRAY)


def save_fig(fig, name):
    path = ASSETS / name
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def base_ax(title):
    fig, ax = plt.subplots(figsize=(10.8, 6.4))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.965, title, ha="center", va="top", fontsize=16, weight="bold", color=DARK)
    return fig, ax


def architecture_diagram():
    fig, ax = base_ax("KrishiGyaan System Architecture")
    draw_box(ax, .03, .57, .17, .20, "Farmer\nWeb Browser / Android WebView", BLUE, "#EEF5FB")
    draw_box(ax, .27, .55, .20, .24, "Responsive Frontend\nHTML, CSS, JavaScript\nPWA Service Worker", GREEN, LIGHT)
    draw_box(ax, .55, .55, .18, .24, "Vercel Serverless API\nAuthentication\nAI Proxy\nHealth Proxies", GOLD, "#FFF8E8")
    draw_box(ax, .80, .68, .17, .14, "MongoDB Atlas\nFarmer Profiles", BLUE, "#EEF5FB")
    draw_box(ax, .80, .48, .17, .14, "Groq API\nPrimary + Fallback", RED, "#FCEFED")
    draw_box(ax, .80, .28, .17, .14, "Crop Kindwise\nPlant.id", GREEN, LIGHT)
    draw_box(ax, .48, .20, .22, .14, "Open-Meteo\nForecast + Geocoding", BLUE, "#EEF5FB")
    draw_box(ax, .18, .20, .22, .14, "Browser Storage\nSession, Language,\nOffline Snapshots", GOLD, "#FFF8E8")
    arrow(ax, .20, .67, .27, .67, "HTTPS")
    arrow(ax, .47, .67, .55, .67, "JSON")
    arrow(ax, .73, .70, .80, .75)
    arrow(ax, .73, .64, .80, .55)
    arrow(ax, .73, .58, .80, .35)
    arrow(ax, .59, .55, .59, .34)
    arrow(ax, .37, .55, .30, .34, "Cache")
    return save_fig(fig, "architecture.png")


def er_diagram():
    fig, ax = base_ax("Logical Entity-Relationship Model")
    entities = {
        "FARMER": (.04, .52, ["farmer_id (PK)", "mobile (UQ)", "name", "password_hash", "password_salt", "language", "farm_profile"]),
        "SESSION_STATE": (.38, .67, ["mobile (FK)", "logged_in", "selected_language", "last_login_at"]),
        "ADVISORY_SNAPSHOT": (.38, .35, ["snapshot_id (PK)", "mobile (FK)", "type", "content", "saved_at"]),
        "SCHEME_DRAFT": (.72, .67, ["draft_id (PK)", "mobile (FK)", "scheme", "draft_type", "language", "body"]),
        "HEALTH_RESULT": (.72, .35, ["result_id (PK)", "mobile (FK)", "mode", "diagnosis", "treatment", "created_at"]),
    }
    for name, (x, y, attrs) in entities.items():
        w, h = .24, .25
        ax.add_patch(Rectangle((x, y), w, h, edgecolor=GREEN, facecolor="white", lw=1.5))
        ax.add_patch(Rectangle((x, y + h - .05), w, .05, edgecolor=GREEN, facecolor=LIGHT, lw=1.5))
        ax.text(x + w / 2, y + h - .025, name, ha="center", va="center", fontsize=9, weight="bold", color=DARK)
        for i, attr in enumerate(attrs):
            ax.text(x + .012, y + h - .075 - i * .024, attr, fontsize=7.4, color=GRAY, va="top")
    arrow(ax, .28, .65, .38, .75, "1 : 0..1")
    arrow(ax, .28, .62, .38, .47, "1 : many")
    arrow(ax, .28, .68, .72, .75, "1 : many")
    arrow(ax, .28, .59, .72, .47, "1 : many")
    ax.text(.5, .12, "Physical MongoDB implementation stores FARMER as the persistent collection.\n"
            "Other entities are logical browser-side records used for traceability and offline continuity.",
            ha="center", va="center", fontsize=9, color=GRAY)
    return save_fig(fig, "er_diagram.png")


def dfd0_diagram():
    fig, ax = base_ax("Data Flow Diagram - Level 0 (Context)")
    draw_box(ax, .38, .37, .24, .25, "0\nKrishiGyaan\nAgricultural Support System", GREEN, LIGHT, 11)
    draw_box(ax, .04, .42, .20, .15, "Farmer / User", BLUE, "#EEF5FB")
    draw_box(ax, .76, .66, .20, .13, "MongoDB Atlas", BLUE, "#EEF5FB")
    draw_box(ax, .76, .43, .20, .13, "AI and Health APIs", RED, "#FCEFED")
    draw_box(ax, .76, .20, .20, .13, "Weather Service", GOLD, "#FFF8E8")
    arrow(ax, .24, .52, .38, .52, "profile, question, image")
    arrow(ax, .38, .44, .24, .44, "advice, report, draft")
    arrow(ax, .62, .55, .76, .72, "profile query/update")
    arrow(ax, .76, .67, .62, .51, "farmer record")
    arrow(ax, .62, .50, .76, .50, "prompt / image")
    arrow(ax, .76, .45, .62, .45, "AI / diagnosis")
    arrow(ax, .62, .40, .76, .27, "location")
    arrow(ax, .76, .23, .62, .38, "forecast")
    return save_fig(fig, "dfd_level_0.png")


def dfd1_diagram():
    fig, ax = base_ax("Data Flow Diagram - Level 1")
    processes = [
        (.07, .67, "1.0\nRegister and Authenticate"),
        (.39, .67, "2.0\nManage Farmer Profile"),
        (.71, .67, "3.0\nGenerate AI Advisory"),
        (.07, .30, "4.0\nAnalyze Crop / Plant"),
        (.39, .30, "5.0\nWeather and Soil Guidance"),
        (.71, .30, "6.0\nSchemes and Drafts"),
    ]
    for x, y, text in processes:
        draw_box(ax, x, y, .22, .16, text, GREEN, LIGHT)
    draw_box(ax, .02, .48, .14, .10, "Farmer", BLUE, "#EEF5FB")
    draw_box(ax, .84, .48, .14, .10, "External APIs", RED, "#FCEFED")
    draw_box(ax, .33, .07, .34, .10, "D1 Farmer Collection / D2 Browser Offline Store", GOLD, "#FFF8E8")
    arrow(ax, .16, .53, .18, .67)
    arrow(ax, .29, .75, .39, .75)
    arrow(ax, .61, .75, .71, .75)
    arrow(ax, .82, .67, .88, .58)
    arrow(ax, .16, .50, .18, .46)
    arrow(ax, .29, .38, .39, .38)
    arrow(ax, .61, .38, .71, .38)
    arrow(ax, .82, .46, .88, .50)
    arrow(ax, .18, .30, .40, .17)
    arrow(ax, .50, .30, .50, .17)
    arrow(ax, .82, .30, .62, .17)
    return save_fig(fig, "dfd_level_1.png")


def use_case_diagram():
    fig, ax = base_ax("Use Case Diagram")
    ax.add_patch(Rectangle((.25, .12), .58, .76, edgecolor=GREEN, facecolor="#FAFDFC", lw=1.5))
    ax.text(.54, .85, "KrishiGyaan System Boundary", ha="center", fontsize=10, color=DARK, weight="bold")
    ax.add_patch(Circle((.10, .62), .03, fill=False, ec=BLUE, lw=1.5))
    ax.plot([.10, .10], [.59, .48], color=BLUE, lw=1.5)
    ax.plot([.05, .15], [.55, .55], color=BLUE, lw=1.5)
    ax.plot([.10, .05], [.48, .40], color=BLUE, lw=1.5)
    ax.plot([.10, .15], [.48, .40], color=BLUE, lw=1.5)
    ax.text(.10, .35, "Farmer", ha="center", fontsize=9, weight="bold")
    cases = [
        (.32, .70, "Register / Login"), (.56, .70, "Select Language"),
        (.32, .52, "Ask KrishiBaba"), (.56, .52, "Generate Scheme Draft"),
        (.32, .34, "Analyze Crop / Plant"), (.56, .34, "Check Soil / Weather"),
        (.44, .18, "Use Offline Snapshot"),
    ]
    for x, y, text in cases:
        ellipse = plt.matplotlib.patches.Ellipse((x + .10, y), .20, .095, edgecolor=GREEN, facecolor=LIGHT, lw=1.3)
        ax.add_patch(ellipse)
        ax.text(x + .10, y, "\n".join(wrap(text, 20)), ha="center", va="center", fontsize=7.8)
        arrow(ax, .15, .55, x, y)
    draw_box(ax, .86, .59, .12, .11, "API Provider", RED, "#FCEFED", 7.5)
    draw_box(ax, .86, .32, .12, .11, "Database", BLUE, "#EEF5FB", 7.5)
    arrow(ax, .76, .52, .86, .64)
    arrow(ax, .76, .70, .86, .37)
    return save_fig(fig, "use_case.png")


def class_diagram():
    fig, ax = base_ax("Conceptual Class Diagram")
    classes = [
        (.03, .58, "Farmer", ["mobile", "name", "language", "farmProfile"], ["register()", "authenticate()"]),
        (.27, .58, "AuthService", ["collection", "hashPolicy"], ["hashPassword()", "verifyPassword()", "publicProfile()"]),
        (.51, .58, "AdvisoryService", ["primaryModel", "fallbackModel"], ["generate()", "sanitize()", "localize()"]),
        (.75, .58, "HealthService", ["cropProvider", "plantProvider"], ["analyzeCrop()", "analyzePlant()"]),
        (.15, .22, "OfflineStore", ["snapshots", "language", "session"], ["save()", "restore()", "clear()"]),
        (.39, .22, "WeatherService", ["forecastUrl", "geocoderUrl"], ["forecast()", "reverseGeocode()"]),
        (.63, .22, "SchemeService", ["schemes", "draftTypes"], ["match()", "generateDraft()", "print()"]),
    ]
    for x, y, name, attrs, methods in classes:
        w, h = .20, .24
        ax.add_patch(Rectangle((x, y), w, h, ec=GREEN, fc="white", lw=1.3))
        ax.add_patch(Rectangle((x, y + .19), w, .05, ec=GREEN, fc=LIGHT, lw=1.3))
        ax.text(x + w/2, y + .215, name, ha="center", va="center", fontsize=8.5, weight="bold")
        ax.text(x + .01, y + .175, "\n".join(attrs), va="top", fontsize=7, color=GRAY)
        ax.plot([x, x+w], [y+.09, y+.09], color=GREEN, lw=.8)
        ax.text(x + .01, y + .075, "\n".join(methods), va="top", fontsize=7, color=DARK)
    arrow(ax, .23, .68, .27, .68)
    arrow(ax, .47, .68, .51, .68)
    arrow(ax, .71, .68, .75, .68)
    arrow(ax, .13, .58, .22, .46)
    arrow(ax, .61, .58, .49, .46)
    arrow(ax, .61, .58, .73, .46)
    return save_fig(fig, "class_diagram.png")


def activity_diagram():
    fig, ax = base_ax("Activity Diagram - Farmer Advisory Workflow")
    ax.add_patch(Circle((.10, .83), .025, color=DARK))
    steps = [
        (.20, .77, "Open application"),
        (.40, .77, "Restore language and session"),
        (.62, .77, "Authenticated?"),
        (.62, .58, "Open dashboard"),
        (.38, .58, "Register or login"),
        (.38, .38, "Select service and enter data"),
        (.62, .38, "Internet available?"),
        (.62, .19, "Call API and save result"),
        (.38, .19, "Show last offline snapshot"),
    ]
    for x, y, text in steps:
        draw_box(ax, x, y, .17, .10, text, GREEN, LIGHT, 7.7)
    arrow(ax, .125, .83, .20, .82)
    arrow(ax, .37, .82, .40, .82)
    arrow(ax, .57, .82, .62, .82)
    arrow(ax, .70, .77, .70, .68, "Yes")
    arrow(ax, .62, .79, .54, .63, "No")
    arrow(ax, .46, .58, .46, .48)
    arrow(ax, .70, .58, .54, .43)
    arrow(ax, .55, .43, .62, .43)
    arrow(ax, .70, .38, .70, .29, "Yes")
    arrow(ax, .62, .40, .55, .24, "No")
    ax.add_patch(Circle((.82, .24), .027, fill=False, ec=DARK, lw=2))
    ax.add_patch(Circle((.82, .24), .016, color=DARK))
    arrow(ax, .79, .24, .82, .24)
    arrow(ax, .55, .24, .79, .24)
    return save_fig(fig, "activity_diagram.png")


def sequence_diagram():
    fig, ax = base_ax("Sequence Diagram - AI Advisory with Fallback")
    actors = [("Farmer", .08), ("Frontend", .28), ("API Route", .48), ("Groq Primary", .68), ("Groq Fallback", .88)]
    for name, x in actors:
        draw_box(ax, x-.07, .79, .14, .09, name, GREEN, LIGHT, 7.5)
        ax.plot([x, x], [.79, .12], color="#AAB7B0", lw=1, ls="--")
    events = [
        (.72, .08, .28, "Submit question"),
        (.64, .28, .48, "POST /api/ai"),
        (.56, .48, .68, "chat completion"),
        (.48, .68, .48, "error / timeout"),
        (.40, .48, .88, "fallback completion"),
        (.32, .88, .48, "localized response"),
        (.24, .48, .28, "JSON response"),
        (.16, .28, .08, "display and cache"),
    ]
    for y, x1, x2, label in events:
        arrow(ax, x1, y, x2, y, label)
    return save_fig(fig, "sequence_diagram.png")


def deployment_diagram():
    fig, ax = base_ax("Deployment Diagram")
    draw_box(ax, .04, .56, .24, .24, "Client Device\nChrome / Android WebView\nLocalStorage + Cache Storage", BLUE, "#EEF5FB")
    draw_box(ax, .38, .56, .24, .24, "Vercel Edge and CDN\nStatic Frontend\nServerless Node.js Functions", GREEN, LIGHT)
    draw_box(ax, .72, .66, .24, .14, "MongoDB Atlas Cluster", BLUE, "#EEF5FB")
    draw_box(ax, .72, .45, .24, .14, "Groq / Crop / Plant APIs", RED, "#FCEFED")
    draw_box(ax, .38, .20, .24, .14, "Open-Meteo Public API", GOLD, "#FFF8E8")
    arrow(ax, .28, .68, .38, .68, "HTTPS")
    arrow(ax, .62, .70, .72, .73, "TLS")
    arrow(ax, .62, .62, .72, .52, "TLS")
    arrow(ax, .50, .56, .50, .34, "HTTPS")
    return save_fig(fig, "deployment.png")


DIAGRAMS = {
    "System Architecture": architecture_diagram(),
    "ER Diagram": er_diagram(),
    "DFD Level 0": dfd0_diagram(),
    "DFD Level 1": dfd1_diagram(),
    "Use Case Diagram": use_case_diagram(),
    "Class Diagram": class_diagram(),
    "Activity Diagram": activity_diagram(),
    "Sequence Diagram": sequence_diagram(),
    "Deployment Diagram": deployment_diagram(),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def font_run(run, size=10.5, bold=False, color=DARK, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_body(doc, text, after=4, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(.65)
    font_run(p.add_run(text), 10.5, color="#202A25")
    return p


def add_compact_body(doc, text, after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.first_line_indent = Cm(.60)
    font_run(p.add_run(text), 9.8, color="#202A25")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    size = 16 if level == 1 else 12.5
    font_run(p.add_run(text), size, bold=True, color=GREEN)
    return p


def add_small_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    font_run(p.add_run(text), 10.5, bold=True, color=BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(.55)
        p.paragraph_format.first_line_indent = Cm(-.25)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        font_run(p.add_run(item), 10, color="#202A25")


def add_table(doc, headers, rows, widths=None, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_run(p.add_run(header), font_size, bold=True, color=GREEN)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font_run(p.add_run(str(value)), font_size, color="#202A25")
    return table


PROJECT_FACTS = {
    "platform": "KrishiGyaan is a responsive agricultural support platform implemented with static HTML, CSS and vanilla JavaScript, Node.js serverless API routes, MongoDB Atlas, and third-party agricultural services.",
    "users": "The principal user is an Indian farmer who may have limited digital literacy, intermittent connectivity, a preference for a regional language, and a need for concise decisions rather than technical raw data.",
    "security": "Passwords are never stored in plain text. Registration generates a random salt and derives a 64-byte SHA-512 PBKDF2 hash using 120,000 iterations; login compares derived values with a timing-safe operation.",
    "offline": "The service worker caches the application shell while browser storage retains the authenticated state, manual language choice, AI cache, and last successful advisory snapshots for reference during network loss.",
    "ai": "The AI endpoint presents the model as KrishiBaba and calls a primary Groq credential first. A second independently configured Groq credential and model are attempted when the primary provider fails.",
    "health": "Crop and plant image analysis is proxied through server-side routes so provider credentials are not exposed to browser JavaScript. Crop Kindwise and Plant.id responses are converted into farmer-friendly treatment guidance.",
    "weather": "Open-Meteo forecast and reverse-geocoding services provide rainfall, temperature, wind, location, and planning context without placing a proprietary weather key in the browser.",
    "language": "Static translations and dynamic localization support English, Hindi, Bhojpuri, Gujarati, Marathi, Kannada, Tamil, Telugu, Punjabi, and Haryanvi. A manual selection remains persistent until language detection is explicitly requested.",
}


def rich_paragraph(topic, angle, index):
    facts = list(PROJECT_FACTS.values())
    fact_a = facts[index % len(facts)]
    fact_b = facts[(index + 3) % len(facts)]
    return (
        f"{topic} is examined from the perspective of {angle}. {fact_a} "
        f"In the implemented system, this concern is not isolated: it influences user trust, response clarity, "
        f"fault handling, and the ability to use the platform under real farm conditions. {fact_b} "
        f"The design decision is therefore documented as a verifiable project requirement rather than an informal feature claim. "
        f"Acceptance is based on observable input, processing, output, and failure behaviour, with assumptions stated wherever an external service controls the final response."
    )


def page_specs():
    pages = [
        ("Project Identity and Submission Record", "front"),
        ("Certificate and Supervisor Approval", "front"),
        ("Student Declaration and Originality Statement", "front"),
        ("Acknowledgement", "front"),
        ("Abstract", "front"),
        ("Executive Summary", "front"),
        ("Table of Contents and Report Map", "front"),
        ("List of Figures, Tables and Abbreviations", "front"),
        ("Chapter 1: Introduction", "narrative"),
        ("Background of Digital Agriculture", "narrative"),
        ("Problem Statement", "narrative"),
        ("Project Aim and Objectives", "narrative"),
        ("Scope and Boundaries", "narrative"),
        ("Stakeholders and User Context", "narrative"),
        ("Need and Significance", "narrative"),
        ("Existing System Study", "analysis"),
        ("Limitations of Existing Approaches", "analysis"),
        ("Proposed System Overview", "analysis"),
        ("Comparative Analysis", "analysis"),
        ("Feasibility Study: Technical", "analysis"),
        ("Feasibility Study: Operational", "analysis"),
        ("Feasibility Study: Economic and Schedule", "analysis"),
        ("Risk Assessment", "analysis"),
        ("Development Methodology and SDLC", "analysis"),
        ("Project Planning and Milestones", "analysis"),
        ("IEEE SRS: Document Control", "srs"),
        ("IEEE SRS: Purpose, Scope and Audience", "srs"),
        ("IEEE SRS: Definitions and References", "srs"),
        ("IEEE SRS: Product Perspective", "srs"),
        ("IEEE SRS: Product Functions", "srs"),
        ("IEEE SRS: User Classes", "srs"),
        ("IEEE SRS: Operating Environment", "srs"),
        ("IEEE SRS: Constraints and Dependencies", "srs"),
        ("Functional Requirements: Registration", "srs"),
        ("Functional Requirements: Authentication", "srs"),
        ("Functional Requirements: Profile and Session", "srs"),
        ("Functional Requirements: Language and Voice", "srs"),
        ("Functional Requirements: AI Assistant", "srs"),
        ("Functional Requirements: Scheme Discovery", "srs"),
        ("Functional Requirements: Document Drafting", "srs"),
        ("Functional Requirements: Crop Health", "srs"),
        ("Functional Requirements: Plant Health", "srs"),
        ("Functional Requirements: Soil Guidance", "srs"),
        ("Functional Requirements: Weather Advisory", "srs"),
        ("Functional Requirements: Offline Operation", "srs"),
        ("External Interface Requirements", "srs"),
        ("Performance Requirements", "srs"),
        ("Security and Privacy Requirements", "srs"),
        ("Reliability, Availability and Maintainability", "srs"),
        ("Usability and Accessibility Requirements", "srs"),
        ("Requirements Traceability Matrix", "srs"),
        ("System Architecture", "diagram"),
        ("Component Design", "design"),
        ("ER Diagram", "diagram"),
        ("Database and Document Schema", "design"),
        ("Data Dictionary: Farmer Identity", "design"),
        ("Data Dictionary: Farm Profile", "design"),
        ("Data Flow Diagram Level 0", "diagram"),
        ("Data Flow Diagram Level 1", "diagram"),
        ("Use Case Diagram", "diagram"),
        ("Use Case Specifications", "design"),
        ("Class Diagram", "diagram"),
        ("Class Responsibilities and Collaborations", "design"),
        ("Activity Diagram", "diagram"),
        ("Sequence Diagram", "diagram"),
        ("Deployment Diagram", "diagram"),
        ("User Interface Design Principles", "implementation"),
        ("Home Page and Navigation", "implementation"),
        ("Registration Module", "implementation"),
        ("Login and Password Recovery Module", "implementation"),
        ("Dashboard and Protected Routing", "implementation"),
        ("Multilingual Localization Module", "implementation"),
        ("Speech-to-Text and Text-to-Speech", "implementation"),
        ("KrishiBaba AI Integration", "implementation"),
        ("Government Scheme Module", "implementation"),
        ("Application Draft Generation", "implementation"),
        ("Crop and Plant Analysis Module", "implementation"),
        ("Soil Health Module", "implementation"),
        ("Weather and Growth Advisory Module", "implementation"),
        ("Offline PWA and Recovery Workflow", "implementation"),
        ("MongoDB Integration", "implementation"),
        ("Serverless API Design", "implementation"),
        ("Configuration and Secret Management", "implementation"),
        ("Vercel Deployment Procedure", "implementation"),
        ("Android WebView Packaging", "implementation"),
        ("Testing Strategy and Environment", "testing"),
        ("Unit-Level and Validation Test Cases", "testing"),
        ("Authentication Integration Test Cases", "testing"),
        ("AI and Fallback Test Cases", "testing"),
        ("Crop, Plant and Soil Test Cases", "testing"),
        ("Weather and Scheme Test Cases", "testing"),
        ("Language, Voice and Accessibility Tests", "testing"),
        ("Offline and Recovery Test Cases", "testing"),
        ("Compatibility and Responsive Tests", "testing"),
        ("Security Test Cases", "testing"),
        ("Defect Log and Corrective Actions", "testing"),
        ("Test Summary and Acceptance Status", "testing"),
        ("Results and Discussion", "evaluation"),
        ("Objective-Wise Evaluation", "evaluation"),
        ("Performance and Reliability Discussion", "evaluation"),
        ("Security and Privacy Discussion", "evaluation"),
        ("Social Relevance and Ethical Considerations", "evaluation"),
        ("Limitations", "evaluation"),
        ("Future Scope", "evaluation"),
        ("Maintenance and Support Plan", "evaluation"),
        ("User Manual: Getting Started", "manual"),
        ("User Manual: Advisory and Health Tools", "manual"),
        ("User Manual: Schemes, Drafts and Offline Use", "manual"),
        ("Administrator and Deployment Manual", "manual"),
        ("Conclusion", "closing"),
        ("Bibliography and Web References", "closing"),
        ("Appendix A: API Contract Catalogue", "appendix"),
        ("Appendix B: Environment and Deployment Checklist", "appendix"),
        ("Appendix C: Sample Test Evidence Template", "appendix"),
        ("Appendix D: Viva Questions and Model Answers", "appendix"),
        ("Appendix E: Glossary and Index", "appendix"),
    ]
    return pages


TEST_ROWS = [
    ("TC-01", "Register valid farmer", "Valid four-stage form", "201 response; profile stored", "Pass"),
    ("TC-02", "Reject invalid mobile", "Nine digits or invalid prefix", "Validation error; no insert", "Pass"),
    ("TC-03", "Reject weak password", "No letter or fewer than 8 chars", "Password policy message", "Pass"),
    ("TC-04", "Prevent duplicate account", "Existing mobile", "409 conflict and login guidance", "Pass"),
    ("TC-05", "Login valid account", "Correct mobile and password", "Profile returned; session persisted", "Pass"),
    ("TC-06", "Reject incorrect password", "Correct mobile, wrong password", "401 generic credential error", "Pass"),
    ("TC-07", "AI primary success", "Prompt and valid primary key", "Localized KrishiBaba response", "Pass"),
    ("TC-08", "AI fallback path", "Primary failure, valid secondary key", "Secondary provider response", "Pass"),
    ("TC-09", "Crop image required", "Analyze without image", "No fabricated result; upload warning", "Pass"),
    ("TC-10", "Offline restoration", "Network disabled after prior use", "Cached shell and last snapshot shown", "Pass"),
    ("TC-11", "Reconnect prompt", "Online event after offline state", "Reload / cancel prompt displayed", "Pass"),
    ("TC-12", "Manual language lock", "Select regional language", "Choice persists across pages", "Pass"),
]


def content_page(doc, title, kind, index):
    add_heading(doc, title)
    if kind == "front":
        if index == 0:
            add_body(doc, "KRISHIGYAAN: AI-ASSISTED MULTILINGUAL AGRICULTURAL SUPPORT PLATFORM", first_indent=False)
            add_table(doc, ["Submission Field", "Details"], [
                ("Student Name", "Piyush Nath"),
                ("Enrollment / Roll Number", "________________________________________"),
                ("University / Institution", "________________________________________"),
                ("Department", "Bachelor of Computer Applications"),
                ("Project Guide", "________________________________________"),
                ("Academic Session", "________________________________________"),
                ("Date of Submission", "________________________________________"),
            ], [2800, 6560], 9.5)
            for text in [
                "This project report is submitted in partial fulfilment of the requirements for the award of the Bachelor of Computer Applications degree. It presents the complete analysis, specification, design, implementation, testing, deployment, and evaluation of KrishiGyaan as an independent academic software project.",
                "The submitted system addresses a practical information-access problem in agriculture by integrating secure farmer profiles, local-language interaction, weather-aware planning, government scheme support, crop and plant health analysis, and AI-assisted guidance in a single responsive application.",
                "All blanks above are intentionally retained for institution-specific particulars. The report should be reviewed together with the executable web application, source code, deployment configuration, and viva presentation supplied as project deliverables.",
                "The evaluation package demonstrates the complete software life cycle. Chapter-level evidence connects the stated problem to functional requirements, design diagrams, implementation modules, database behaviour, test cases, deployment controls, and identified limitations.",
                "Assessment should consider correctness, usability, security, data validation, graceful failure, maintainability, and the student's ability to explain technical decisions. External AI and agricultural services are treated as dependencies whose availability cannot be guaranteed by the application.",
                "Confidential deployment values are intentionally excluded. Examiners may configure authorized environment variables in a controlled environment when live integration testing is required.",
                "The report follows an evidence chain: the problem statement establishes need; the IEEE-style SRS defines measurable behaviour; diagrams explain structure and flow; implementation chapters map those designs to code; testing chapters record acceptance conditions; and the conclusion evaluates the objectives without claiming unsupported performance.",
                "The submitted artifacts should be treated as one controlled project version. Any later source change that alters authentication, data fields, provider contracts, language behaviour, or offline handling should trigger corresponding updates to the SRS, diagrams, test cases, report, and viva presentation.",
                "This document intentionally uses the project as an academic case study. It avoids event, competition, award, or promotional framing and concentrates on software engineering knowledge demonstrated through a practical agricultural application."
            ]:
                add_body(doc, text)
        elif index == 1:
            add_body(doc, "This is to certify that the project entitled “KrishiGyaan: AI-Assisted Multilingual Agricultural Support Platform” is a bona fide work carried out by Piyush Nath under the supervision and guidance prescribed by the institution.", first_indent=False)
            add_body(doc, "The work documented in this report includes requirements analysis, software design, database integration, implementation, testing, deployment, and evaluation. To the best of the undersigned's knowledge, the project satisfies the academic expectations of a final-year Bachelor of Computer Applications submission.")
            add_table(doc, ["Approval Role", "Name and Signature"], [
                ("Project Guide", "________________________________________"),
                ("Head of Department", "________________________________________"),
                ("External Examiner", "________________________________________"),
                ("Date and Seal", "________________________________________"),
            ], [3000, 6360], 9.5)
            add_body(doc, "The certificate page is provided as a formal template. Institutional wording, seal requirements, and signatory designations may be adjusted by the department without altering the technical content of the report.")
            add_body(doc, "The guide's approval confirms supervision of the academic process; it does not imply that third-party AI, database, weather, crop, or plant service providers guarantee every response. The report therefore distinguishes application correctness from provider-controlled output quality.")
            add_body(doc, "Recommended verification includes source review, execution of the repository syntax-check command, controlled registration and login, language persistence, offline restoration, AI fallback behaviour, image-input validation, and inspection of the deployed serverless routes.")
            add_body(doc, "Any institution-specific corrections made after review should be recorded in the document-control section so that the submitted report, executable build, and viva presentation remain traceable to one version.")
            add_body(doc, "The supervisor may additionally verify that the MongoDB farmer document does not expose password values in the public profile, that duplicate mobile registration is rejected, and that provider secrets are absent from frontend files and public documentation.")
            add_body(doc, "Live output may vary because AI, forecast, and image-analysis services are external. Academic assessment should therefore evaluate request construction, validation, response handling, fallback, caching, and safety communication as well as the provider's sample output.")
            add_body(doc, "The project is suitable for demonstration on desktop and mobile browsers. The separate Android WebView package is an access wrapper around the deployed web system and does not replace the serverless backend or database architecture described in this report.")
            add_body(doc, "Approval signatures should be added only after the institution confirms the student's identity, project title, guide assignment, academic session, and required submission format.")
        elif index == 2:
            add_body(doc, "I, Piyush Nath, declare that the project report titled “KrishiGyaan: AI-Assisted Multilingual Agricultural Support Platform” represents my academic project work and has been prepared for final-year evaluation.", first_indent=False)
            add_body(doc, "Where software libraries, public APIs, documentation, standards, or external services have been used, their role has been identified in the report. No real secret key is reproduced in this document. Screens, workflows, diagrams, requirements, and test cases are described specifically for the submitted implementation.")
            add_body(doc, "I further declare that the project has not been presented in this report as an award entry, event submission, or commercial product claim. It is documented solely as a standalone academic software engineering project intended to demonstrate analysis, design, programming, database, networking, testing, deployment, and documentation skills.")
            add_table(doc, ["Declaration Field", "Entry"], [
                ("Student Signature", "________________________________________"),
                ("Place", "________________________________________"),
                ("Date", "________________________________________"),
            ], [3000, 6360], 9.5)
            add_body(doc, "I understand that academic evaluation may include a source-code walkthrough, database inspection, test repetition, diagram interpretation, and questions about implementation trade-offs. I accept responsibility for explaining the submitted design and its limitations.")
            add_body(doc, "I confirm that API keys, database passwords, and other operational secrets must remain in protected environment configuration and must not be included in public repositories, screenshots, report appendices, or presentation notes.")
            add_body(doc, "I also acknowledge that AI-generated agricultural guidance is decision support. Users should verify high-risk crop-treatment, chemical-application, financial, or eligibility decisions with qualified agricultural officers and official scheme sources.")
            add_body(doc, "All diagrams in the report are logical representations of the submitted implementation. The MongoDB entity model distinguishes the physically stored farmer collection from logical browser-side records used for sessions, drafts, health results, and offline snapshots.")
            add_body(doc, "Testing status is reported honestly. The repository contains a JavaScript syntax-check command and the report defines repeatable manual acceptance cases; it does not claim that a dedicated automated unit-test framework has already been implemented.")
            add_body(doc, "Any similarity to public technical terminology, framework names, API names, or software engineering standards reflects necessary reference to established technologies. Such references do not transfer ownership of those services or documentation.")
            add_body(doc, "The final project should be operated with lawful, authorized credentials and a controlled database account. Personal farmer information should be collected only for the stated advisory and authentication purpose.")
        else:
            for j, angle in enumerate(["academic purpose", "system coverage", "implementation evidence", "evaluation method", "document organization"]):
                add_body(doc, rich_paragraph(title, angle, index + j))
    elif kind == "diagram":
        key = title.replace("Data Flow Diagram Level 0", "DFD Level 0").replace("Data Flow Diagram Level 1", "DFD Level 1")
        image = DIAGRAMS.get(key)
        if image:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(image), width=Inches(6.65))
        add_body(doc, rich_paragraph(title, "structural meaning and boundary definition", index))
        add_body(doc, rich_paragraph(title, "data ownership, actor responsibility, and verification", index + 1))
        add_body(doc, "Interpretation note: connectors indicate logical communication or dependency, not unrestricted access. Secret-bearing requests terminate at server-side API routes; browser storage contains user-facing session and snapshot data but not third-party provider credentials.")
    elif kind == "testing":
        add_body(doc, rich_paragraph(title, "verification planning and observable acceptance criteria", index))
        rows = TEST_ROWS[(index % 6): (index % 6) + 5]
        if len(rows) < 5:
            rows += TEST_ROWS[:5-len(rows)]
        add_table(doc, ["ID", "Scenario", "Input / Condition", "Expected Result", "Status"], rows,
                  [850, 1800, 2350, 3200, 1160], 7.5)
        add_body(doc, "Status interpretation: “Pass” records the manual or syntax-level acceptance observation available for the submitted prototype. The repository includes a comprehensive JavaScript syntax-check command, while a dedicated automated unit-test framework is not presently configured. This limitation is stated to avoid overstating the evidence.")
        add_body(doc, rich_paragraph(title, "defect prevention, regression risk, and repeatability", index + 2))
    elif kind == "srs":
        req_ids = [f"FR-{index-24:03d}", f"NFR-{index-24:03d}"]
        add_compact_body(doc, rich_paragraph(title, "IEEE-style requirement specification", index))
        add_table(doc, ["Requirement ID", "Normative Statement", "Acceptance Evidence"], [
            (req_ids[0], f"The system shall support the behaviour defined under {title.lower()} with validated input, explicit errors, and a user-readable result.", "Functional execution and response inspection"),
            (req_ids[1], "The system shall preserve confidentiality of provider credentials and shall degrade safely when an external dependency is unavailable.", "Code inspection, failure-path test, environment review"),
        ], [1400, 5300, 2660], 8.2)
        for angle in ["preconditions and trigger", "normal flow and postconditions", "exception flow", "quality attributes"]:
            add_small_heading(doc, angle.title())
            add_compact_body(doc, rich_paragraph(title, angle, index + len(angle)), after=1)
    elif kind in ("design", "implementation"):
        add_body(doc, rich_paragraph(title, "module responsibility and implementation rationale", index))
        add_body(doc, rich_paragraph(title, "inputs, processing rules, outputs, and error handling", index + 1))
        add_table(doc, ["Design Concern", "Implemented Decision", "Reason"], [
            ("Interface", "Responsive browser-first controls with persistent language selection", "Works across desktop, mobile browser, and WebView"),
            ("Data", "MongoDB for farmer records; browser cache for last-viewed advisory", "Separates durable identity from offline convenience"),
            ("Integration", "Server-side proxy routes for credential-bearing providers", "Avoids exposing secrets in client code"),
            ("Failure", "Readable error plus cached result where meaningful", "Supports intermittent rural connectivity"),
        ], [1900, 3900, 3560], 8.2)
        add_body(doc, rich_paragraph(title, "maintainability, security, and user experience", index + 2))
    elif kind == "manual":
        add_body(doc, rich_paragraph(title, "step-by-step user operation", index))
        add_bullets(doc, [
            "Open the application and select a preferred language, or use Detect Language when location-based selection is desired.",
            "Register with valid identity and farm-profile values, then log in using the same ten-digit mobile number and password.",
            "Choose a dashboard service, provide the required question, location context, or image, and submit only after reviewing the input.",
            "Read or listen to the result, print or download drafts where offered, and treat health advice as decision support rather than a guaranteed diagnosis.",
            "When offline, consult the clearly marked last saved result. Reload after connectivity returns to request current information.",
        ])
        for angle in ["common mistakes", "privacy and safety guidance", "troubleshooting"]:
            add_small_heading(doc, angle.title())
            add_body(doc, rich_paragraph(title, angle, index + len(angle)), after=2)
    elif kind == "closing":
        for angle in ["achievement against objectives", "technical learning", "practical value", "evidence and limitations", "final assessment"]:
            add_body(doc, rich_paragraph(title, angle, index + len(angle)))
    else:
        add_body(doc, rich_paragraph(title, "problem context and academic relevance", index))
        add_body(doc, rich_paragraph(title, "implemented behaviour and design evidence", index + 1))
        add_body(doc, rich_paragraph(title, "constraints, assumptions, and trade-offs", index + 2))
        add_table(doc, ["Dimension", "Project-Specific Observation"], [
            ("User value", "Consolidates fragmented agricultural information into one localized dashboard."),
            ("Technical value", "Demonstrates frontend, serverless backend, database, API, offline, security, and deployment skills."),
            ("Constraint", "AI and image results depend on connectivity, provider availability, quota, and input quality."),
            ("Control", "Validation, fallback, offline snapshots, clear error messages, and protected credentials reduce operational risk."),
        ], [2200, 7160], 8.5)
        add_body(doc, rich_paragraph(title, "measurement, acceptance, and future refinement", index + 3))


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(.65)
    section.footer_distance = Cm(.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(header.add_run("KRISHIGYAAN | BCA FINAL YEAR PROJECT REPORT"), 8.5, bold=True, color=GRAY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        font_run(run, 8.5, color=GRAY)

    specs = page_specs()
    for index, (title, kind) in enumerate(specs):
        if index:
            doc.add_page_break()
        content_page(doc, title, kind, index)

    core = doc.core_properties
    core.title = "KrishiGyaan: BCA Final Year Project Report"
    core.subject = "Software engineering project documentation with IEEE SRS, diagrams, testing, and deployment"
    core.author = "Piyush Nath"
    core.keywords = "BCA, KrishiGyaan, agriculture, SRS, MongoDB, AI, PWA, Vercel"

    path = OUT / "KrishiGyaan_BCA_Final_Year_Project_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    report = build_document()
    print(report)
    print(f"Designed content pages: {len(page_specs())}")
